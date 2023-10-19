import os

from django.conf import settings

from docx.shared import Inches


def data_table(data):
    table_data = data['values'][0]['values']

    matrix = []

    for row in table_data:
        row_values = []
        for cell in row['values']:
            row_values.append(cell['values'])
        matrix.append(row_values)

    return matrix


class HTMLtoDocx:
    def __init__(self, doc):
        self.styles = {
            "p": "paragrafo",
            "h1": "Heading 1",
            "h2": "Heading 2",
            "h3": "Heading 3",
            "h4": "Heading 4",
            "h5": "Heading 5",
            "blockquote": "citação direta",
            "div": "referencias"
        }

        self.OPTIONS = {
            "bold": False,
            "italic": False,
            "underline": False,
        }

        self.sub_document = doc.new_subdoc()

    def render_values(self, paragraph, item, options=None):
        options = options or self.OPTIONS
        for value in item:
            run = paragraph.add_run()
            if not paragraph.style.name.startswith("Heading"):
                run.italic = options["italic"]
                run.bold = options["bold"]
                run.underline = options["underline"]
            if isinstance(value, str):
                run.add_text(value)
            elif value["tag_name"] == "img":
                width = value["attributes"].get("width", "")
                height = value["attributes"].get("height", "")
                source = value["attributes"]["src"]
                try:
                    source = os.path.join(settings.MEDIA_ROOT, source.split('/media/')[1])
                    if width == "":
                        run.add_picture(source)
                    else:
                        width = Inches(16 * 0.393701)
                        # height = Inches(float(height) * 0.0138889)
                        run.add_picture(source, width=width)
                    paragraph.style = 'imagem'
                except Exception:
                    pass
            elif value["tag_name"] == "strong":
                run.bold = True
                new_options = {
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                }
                self.render_values(paragraph, value["values"], new_options)
            elif value["tag_name"] == "em":
                run.italic = True
                new_options = {
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                }
                self.render_values(paragraph, value["values"], new_options)
            elif value["tag_name"] == "span":
                run.underline = True
                new_options = {
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                }
                self.render_values(paragraph, value["values"], new_options)

    def convert_list(self, li, nivel):
        for item in li["values"]:
            paragrafo = None
            if nivel == 1:
                paragrafo = self.sub_document.add_paragraph(style=f"lista{nivel}")
            if len(item["values"]) > 1:
                print(item["values"])
                self.render_values(paragrafo, item["values"][0])
                self.convert_list(item["values"][1], nivel + 1)
            else:
                if paragrafo is None:
                    paragrafo = self.sub_document.add_paragraph(style=f"lista{nivel}")
                    self.render_values(paragrafo, item["values"])
                else:
                    self.render_values(paragrafo, item["values"])

    def convert_table(self, data):
        table = self.sub_document.add_table(
            rows=len(data), cols=len(data[0]), style="Table Grid")

        for row_index, row in enumerate(data):
            row_cells = table.rows[row_index].cells
            for index, cell_data in enumerate(row):
                p = row_cells[index].paragraphs[0]
                p.style = 'Normal'
                if isinstance(cell_data[0], str):
                    self.render_values(p, cell_data)
                else:
                    for index_item, item in enumerate(cell_data):
                        if isinstance(item, str):
                            self.render_values(p, item)
                        elif index_item == 0:
                            self.render_values(p, item["values"])
                        else:
                            paragraph = row_cells[index].add_paragraph()
                            paragraph.style = 'Normal'
                            self.render_values(paragraph, item["values"])

    def convert_default(self, item):
        paragraph = self.sub_document.add_paragraph()
        paragraph.style = self.styles[item['tag_name']]
        self.render_values(paragraph, item['values'])

    def convert(self, items):

        for key in items["keys"]:
            if key["tag_name"] == "ul":
                self.convert_list(key, 1)
            elif key["tag_name"] == "table":
                data = data_table(key)
                self.convert_table(data)
            else:
                self.convert_default(key)

        return self.sub_document
